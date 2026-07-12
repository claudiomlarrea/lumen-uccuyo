"""Generación del Orden del Día en Word (prototipo local)."""

from __future__ import annotations

import re
import unicodedata
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from data.investigacion import parse_puntaje
from data.orden_cs import ordenar_temas_consejo_superior, segmentos_grupo_cs

ROOT = Path(__file__).resolve().parent.parent
LOGO = ROOT / "assets" / "logo_uccuyo.png"
if not LOGO.exists():
    LOGO = ROOT / "assets" / "logo_uccuyo_badge.png"

VERDE = RGBColor(0x04, 0x4A, 0x30)
ROJO = RGBColor(0x8B, 0x1E, 0x2D)


def _slug(texto: str) -> str:
    """Nombre de archivo seguro (ASCII) para evitar descargas sin extensión en Edge/Chrome."""
    nfkd = unicodedata.normalize("NFKD", texto)
    ascii_text = nfkd.encode("ascii", "ignore").decode("ascii")
    limpio = re.sub(r"[^\w-]", "_", ascii_text)
    limpio = re.sub(r"_+", "_", limpio).strip("_")
    return (limpio or "documento")[:50]


TIPOS_CON_DIRECTOR = frozenset(
    {
        "Proyecto de Investigación",
        "Proyecto de Cátedra",
        "Informe Final",
        "Informe de Avance",
    }
)


def puntaje_texto_para_word(raw: Any) -> str | None:
    """Texto para línea «Puntaje: …» del Word (coma decimal, 2 decimales)."""
    if raw in (None, ""):
        return None
    n = parse_puntaje(raw)
    if n is None or n <= 0:
        return None
    x = float(n)
    for _ in range(10):
        if x <= 1000:
            break
        if abs(x - round(x)) >= 1e-4:
            break
        ri = int(round(x))
        if ri % 100 != 0:
            break
        x = x / 100.0
    if x <= 0 or x > 1000:
        return None
    return f"{x:.2f}".replace(".", ",")


def _formato_monto(monto: Any) -> str:
    try:
        valor = int(float(monto))
        return f"${valor:,}".replace(",", ".")
    except (TypeError, ValueError):
        return str(monto or "")


def _linea_director(nombre: str, categoria: str) -> str:
    nombre = str(nombre or "").strip()
    categoria = str(categoria or "").strip()
    if not nombre and not categoria:
        return ""
    if not categoria or categoria.startswith("Seleccionar"):
        return f"   Director: {nombre}\n"
    return f"   Director: {nombre} ({categoria})\n"


def _linea_codirector(nombre: str, categoria: str) -> str:
    nombre = str(nombre or "").strip()
    categoria = str(categoria or "").strip()
    if not nombre and not categoria:
        return ""
    if not categoria or categoria.startswith("Seleccionar"):
        return f"   Codirector: {nombre}\n"
    return f"   Codirector: {nombre} ({categoria})\n"


def _agregar_tema_ci(doc: Document, contador: int, tema: dict[str, Any]) -> int:
    inv = tema.get("investigacion") or {}
    tipo = inv.get("tipo") or tema.get("tipo_actividad", "")
    titulo = inv.get("titulo") or tema.get("actividad", "")
    descripcion = inv.get("descripcion") or tema.get("detalle", "")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1

    run = p.add_run(f"{contador}. {tipo} - {titulo}\n")
    run.bold = True
    run.font.color.rgb = VERDE

    if descripcion:
        p.add_run(f"   Descripción: {descripcion}\n")

    if tipo == "Categorización Docente":
        if inv.get("apellido_nombre_docente"):
            p.add_run(f"   Docente: {inv['apellido_nombre_docente']}\n")
        if inv.get("dni_docente"):
            p.add_run(f"   DNI: {inv['dni_docente']}\n")
    elif tipo in TIPOS_CON_DIRECTOR:
        linea = _linea_director(inv.get("director", ""), inv.get("cat_director", ""))
        if linea:
            p.add_run(linea)
        linea = _linea_codirector(inv.get("codirector", ""), inv.get("cat_codirector", ""))
        if linea:
            p.add_run(linea)

    equipo = str(inv.get("equipo") or "").strip()
    if equipo:
        p.add_run(f"   Equipo: {equipo.replace(chr(10), '; ')}\n")

    ua_word = inv.get("unidades_academicas") or tema.get("unidad_academica", "—")
    p.add_run(f"   Unidad Académica: {ua_word}\n")

    txt_puntaje = puntaje_texto_para_word(inv.get("puntaje"))
    if txt_puntaje:
        p.add_run(f"   Puntaje: {txt_puntaje}\n")

    if inv.get("resolucion_cd"):
        p.add_run(f"   Resolución CD: {inv['resolucion_cd']}\n")
    if inv.get("resolucion_cs"):
        p.add_run(f"   Resolución CS del Proyecto: {inv['resolucion_cs']}\n")
    if inv.get("instituto"):
        p.add_run(f"   Instituto: {inv['instituto']}\n")
    if inv.get("catedra"):
        p.add_run(f"   Cátedra: {inv['catedra']}\n")
    if inv.get("tipo_financiamiento"):
        p.add_run(f"   Financiamiento: {inv['tipo_financiamiento']}\n")
    if inv.get("fuente_financiamiento"):
        p.add_run(f"   Fuente: {inv['fuente_financiamiento']}\n")
    if inv.get("responsable_de_carga"):
        p.add_run(f"   Responsable de carga: {inv['responsable_de_carga']}\n")
    if inv.get("monto_financiamiento"):
        p.add_run(f"   Monto: {_formato_monto(inv['monto_financiamiento'])}\n")
    if inv.get("alumnos"):
        p.add_run(f"   Alumnos: {inv['alumnos']}\n")

    return contador + 1


def _agregar_tema_cs(doc: Document, contador: int, tema: dict[str, Any], *, omitir_unidad: bool) -> int:
    head = doc.add_paragraph()
    rh = head.add_run(f"{contador}. {tema.get('actividad', '(sin título)')}")
    rh.bold = True
    rh.font.size = Pt(11)
    rh.font.color.rgb = VERDE

    lines = [
        f"Sesión: {tema.get('fecha_reunion', '—')}",
        f"Tipo: {tema.get('tipo_actividad', '—')}",
        f"Ámbito: {tema.get('ambito', '—')}",
    ]
    if not omitir_unidad:
        lines.insert(0, f"Unidad: {tema.get('unidad_academica', '—')}")
    if tema.get("elevado_desde_cd"):
        lines.append(f"Elevado desde CD: {tema.get('fecha_cd', '—')}")
    if tema.get("impacta_pei"):
        lines.append(f"Objetivo PEI: {tema.get('objetivo_especifico', '—')}")
    else:
        lines.append("Impacta PEI: No")
    if tema.get("detalle"):
        lines.append(f"Detalle: {tema['detalle']}")
    lines.append(f"ID: {tema.get('id', '—')}")

    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph("")
    return contador + 1


def _renderizar_temas_consejo_superior(
    doc: Document,
    temas: list[dict[str, Any]],
    orden_ua: list[str] | None = None,
) -> None:
    temas_ordenados = ordenar_temas_consejo_superior(temas, orden_ua=orden_ua)
    contador = 1
    for encabezado, items in segmentos_grupo_cs(temas_ordenados):
        if encabezado:
            doc.add_paragraph("")
            h = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_h = h.add_run(encabezado)
            run_h.bold = True
            run_h.font.size = Pt(12)
            run_h.font.color.rgb = RGBColor(0, 102, 204)
        for tema in items:
            contador = _agregar_tema_cs(doc, contador, tema, omitir_unidad=bool(encabezado))


def generar_orden_del_dia(
    temas: list[dict[str, Any]],
    unidad: str,
    sede: str,
    anio: str,
    fecha_reunion: str | None = None,
    organo: str = "Consejo Directivo",
    orden_ua: list[str] | None = None,
) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Encabezado institucional centrado
    if LOGO.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.add_run().add_picture(str(LOGO), width=Cm(2.2))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Universidad Católica de Cuyo")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = VERDE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("LUMEN — Orden del Día Institucional")
    r2.font.size = Pt(11)
    r2.font.color.rgb = ROJO

    doc.add_paragraph("")

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = titulo.add_run("ORDEN DEL DÍA")
    rt.bold = True
    rt.font.size = Pt(16)
    rt.font.color.rgb = VERDE

    subt = doc.add_paragraph()
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = subt.add_run(organo)
    rs.bold = True
    rs.font.size = Pt(12)
    rs.font.color.rgb = ROJO

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{unidad}\n").bold = True
    meta.add_run(f"Sede: {sede}  ·  Año: {anio}\n")
    if fecha_reunion:
        meta.add_run(f"Fecha de reunión: {fecha_reunion}\n")
    elif temas:
        fechas = sorted({t.get("fecha_reunion") for t in temas if t.get("fecha_reunion")})
        if len(fechas) == 1:
            meta.add_run(f"Fecha de reunión: {fechas[0]}\n")
    meta.add_run(f"Generado: {datetime.now():%d/%m/%Y %H:%M}")

    doc.add_paragraph("")
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if organo == "Consejo Superior":
        texto = (
            "Temas para tratamiento en Consejo Superior. "
            "Incluye propuestas elevadas por unidades académicas y cargadas por Rectorado / SGA."
        )
    else:
        texto = f"Temas propuestos para tratamiento en {organo}."
    intro.add_run(texto + " Prototipo LUMEN (no impacta planillas productivas).").italic = True

    doc.add_paragraph("")

    if not temas:
        vacio = doc.add_paragraph("No hay temas cargados para los filtros seleccionados.")
        vacio.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif organo == "Consejo de Investigación":
        contador = 1
        unidad_actual = ""
        for tema in temas:
            unidad_tema = tema.get("unidad_academica", "—")
            if unidad_tema != unidad_actual:
                doc.add_paragraph("")
                h = doc.add_paragraph()
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_h = h.add_run(unidad_tema)
                run_h.bold = True
                run_h.font.size = Pt(12)
                run_h.font.color.rgb = RGBColor(0, 102, 204)
                unidad_actual = unidad_tema
            contador = _agregar_tema_ci(doc, contador, tema)
    elif organo == "Consejo Superior":
        _renderizar_temas_consejo_superior(doc, temas, orden_ua=orden_ua)
    else:
        for i, tema in enumerate(temas, start=1):
            head = doc.add_paragraph()
            rh = head.add_run(f"{i}. {tema.get('actividad', '(sin título)')}")
            rh.bold = True
            rh.font.size = Pt(11)
            rh.font.color.rgb = VERDE

            lines = [
                f"Unidad: {tema.get('unidad_academica', '—')}",
                f"Sesión: {tema.get('fecha_reunion', '—')}",
                f"Tipo: {tema.get('tipo_actividad', '—')}",
                f"Ámbito: {tema.get('ambito', '—')}",
            ]
            if tema.get("elevado_desde_cd"):
                lines.append(f"Elevado desde CD: {tema.get('fecha_cd', '—')}")
            if tema.get("impacta_pei"):
                lines.append(f"Objetivo PEI: {tema.get('objetivo_especifico', '—')}")
            else:
                lines.append("Impacta PEI: No")
            if tema.get("detalle"):
                lines.append(f"Detalle: {tema['detalle']}")
            lines.append(f"ID: {tema.get('id', '—')}")

            for line in lines:
                p = doc.add_paragraph(line)
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(2)

            doc.add_paragraph("")

    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp = pie.add_run("UCCuyo · Testimonium de Lumine · Prototipo LUMEN")
    rp.font.size = Pt(9)
    rp.font.color.rgb = VERDE

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def nombre_archivo_od(
    organo: str,
    unidad: str,
    anio: str,
    fecha_reunion: str | None = None,
    fecha_iso: str | None = None,
) -> str:
    prefijos = {
        "Consejo Superior": "CS",
        "Consejo de Investigación": "CI",
        "Consejo de Extensión": "CE",
        "Consejo Directivo": "CD",
    }
    prefijo = prefijos.get(organo, "OD")
    slug_ua = _slug(unidad) if unidad != "Todas las unidades" else "Institucional"
    if fecha_iso:
        slug_fecha = fecha_iso.replace("-", "")
    elif fecha_reunion:
        slug_fecha = _slug(fecha_reunion)[:24]
    else:
        slug_fecha = anio
    return f"LUMEN_OD_{prefijo}_{slug_ua}_{slug_fecha}.docx"
